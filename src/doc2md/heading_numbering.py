"""Extract heading numbering information from DOCX files."""

from __future__ import annotations

import re
from typing import Dict, List, Tuple
from docx import Document


def extract_heading_numbering_from_toc(docx_path: str) -> Dict[str, str]:
    """
    Extract heading numbering from the Table of Contents in a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping heading text (without numbers) to their numbers
        Example: {"Общие сведения": "1", "Назначение": "1.1", "Подготовка конфигурационных файлов": "4.1.2.1"}
    """
    doc = Document(docx_path)
    numbering_map = {}
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Check if this is a TOC entry with numbering
        if paragraph.style.name.startswith('toc'):
            # Pattern to match numbered TOC entries like "4.1.2.1 Подготовка конфигурационных файлов\t42"
            match = re.match(r'^(\d+(?:\.\d+)*)\s+([^\t]+)(?:\t\d+)?$', text)
            if match:
                number = match.group(1)
                title = match.group(2).strip()
                numbering_map[title] = number
                
    return numbering_map


def extract_heading_structure_from_toc(docx_path: str) -> List[Tuple[int, str, str]]:
    """
    Extract complete heading structure from TOC.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        List of tuples (level, number, title) sorted by document order
        Example: [(1, "1", "Общие сведения"), (2, "1.1", "Назначение"), (4, "4.1.2.1", "Подготовка")]
    """
    doc = Document(docx_path)
    headings = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Check if this is a TOC entry
        if paragraph.style.name.startswith('toc'):
            # Extract level from style name (toc 1, toc 2, etc.)
            level_match = re.search(r'toc (\d+)', paragraph.style.name)
            if not level_match:
                continue
                
            level = int(level_match.group(1))
            
            # Extract number and title
            match = re.match(r'^(\d+(?:\.\d+)*)\s+([^\t]+)(?:\t\d+)?$', text)
            if match:
                number = match.group(1)
                title = match.group(2).strip()
                headings.append((level, number, title))
                
    return headings


def get_heading_number_for_text(text: str, numbering_map: Dict[str, str]) -> str | None:
    """
    Find the heading number for a given text by fuzzy matching against the numbering map.
    
    Args:
        text: The heading text to match
        numbering_map: Dictionary from extract_heading_numbering_from_toc
        
    Returns:
        The heading number if found, None otherwise
    """
    text = text.strip()
    
    # First try exact match
    if text in numbering_map:
        return numbering_map[text]
    
    # Try fuzzy matching - look for text that contains the same words
    text_words = set(text.lower().split())
    best_match = None
    best_score = 0
    
    for toc_title, number in numbering_map.items():
        toc_words = set(toc_title.lower().split())
        
        # Calculate similarity score (intersection over union)
        intersection = len(text_words & toc_words)
        union = len(text_words | toc_words)
        
        if union > 0:
            score = intersection / union
            if score > best_score and score > 0.5:  # At least 50% similarity
                best_score = score
                best_match = number
                
    return best_match


def add_numbering_to_html(html_content: str, docx_path: str) -> str:
    """
    Add heading numbering to HTML content based on DOCX TOC.
    Uses regex to find and replace __RefHeading patterns with proper heading tags.
    
    Args:
        html_content: HTML content from Mammoth conversion
        docx_path: Path to original DOCX file
        
    Returns:
        HTML content with numbered headings
    """
    import re
    
    # Extract numbering information from DOCX
    numbering_map = extract_heading_numbering_from_toc(docx_path)
    heading_structure = extract_heading_structure_from_toc(docx_path)
    
    if not numbering_map:
        return html_content  # No numbering found
    
    # Create a map of titles to their levels and numbers
    title_to_info = {}
    for level, number, title in heading_structure:
        title_to_info[title.lower()] = (level, number, title)
    
    # Use regex to find patterns like: <a id="__RefHeading___X"></a>TextHere
    # This pattern matches anchor + immediate heading text
    pattern = r'<a id="__RefHeading___\d+"></a>([А-ЯЁA-Z][а-яёa-z\s]+?)(?=Программный|[.!?]|[А-ЯЁA-Z]{3,}|<|\s{2,}|$)'
    
    def replace_heading(match):
        heading_text = match.group(1).strip()
        heading_lower = heading_text.lower()
        
        # Find exact match in TOC first
        if heading_lower in title_to_info:
            level, number, original_title = title_to_info[heading_lower]
            return f'<h{level}>{number} {original_title}</h{level}>'
        
        # Try fuzzy matching
        best_match = None
        best_score = 0
        
        text_words = set(heading_lower.split())
        for toc_title_lower, (level, number, original_title) in title_to_info.items():
            toc_words = set(toc_title_lower.split())
            
            if len(text_words) > 0 and len(toc_words) > 0:
                intersection = len(text_words & toc_words)
                union = len(text_words | toc_words)
                score = intersection / union if union > 0 else 0
                
                if score > best_score and score > 0.8:  # High threshold
                    best_score = score
                    best_match = (level, number, original_title)
        
        if best_match:
            level, number, original_title = best_match
            return f'<h{level}>{number} {original_title}</h{level}>'
        
        return match.group(0)  # No change if not found
    
    # Apply replacement
    result = re.sub(pattern, replace_heading, html_content)
    return result


if __name__ == "__main__":
    # Test the functionality
    import sys
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        print("Extracting heading numbering from TOC...")
        numbering_map = extract_heading_numbering_from_toc(docx_path)
        
        print("Found numbering:")
        for title, number in sorted(numbering_map.items(), key=lambda x: x[1]):
            print(f"  {number}: {title}")
        
        print(f"\nTotal headings: {len(numbering_map)}")
        
        print("\nHeading structure:")
        structure = extract_heading_structure_from_toc(docx_path)
        for level, number, title in structure[:20]:  # Show first 20
            indent = "  " * (level - 1)
            print(f"{indent}H{level} {number} {title}")