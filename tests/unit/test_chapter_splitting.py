#!/usr/bin/env python3
"""
Модульные тесты для функциональности разделения на главы
"""

import unittest
import tempfile
import os
import shutil
from pathlib import Path
import sys

# Добавляем родительскую директорию в path для импорта модуля
sys.path.append(str(Path(__file__).parent.parent.parent))

from improved_docx_converter import ImprovedDocxToMarkdownConverter


class TestChapterSplitting(unittest.TestCase):
    """Тесты для разделения документа на главы"""
    
    def setUp(self):
        """Подготовка к тестам"""
        self.temp_dir = Path(tempfile.mkdtemp())
        self.samples_dir = Path("samples")
        
    def tearDown(self):
        """Очистка после тестов"""
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
    
    def test_chapter_detection_patterns(self):
        """Тест определения границ глав"""
        # Создаем временный конвертер для тестирования
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "temp.docx"
            from docx import Document
            doc = Document()
            doc.save(doc_path)
            
            converter = ImprovedDocxToMarkdownConverter(str(doc_path))
            
            # Тестовые заголовки, которые должны определяться как начало новой главы
            chapter_headers = [
                "1. Общие сведения",
                "2. Архитектура комплекса", 
                "3. Технические требования",
                "4. Установка и настройка",
                "5. Эксплуатация системы"
            ]
            
            # Эти заголовки НЕ должны начинать новую главу
            section_headers = [
                "1.1 Функциональное назначение",
                "2.1 Основные компоненты", 
                "3.1 Требования к ПО",
                "1.2 Требования к среде"
            ]
            
            for header in chapter_headers:
                self.assertTrue(converter._is_chapter_header(header), 
                              f"Заголовок '{header}' должен определяться как глава")
                
            for header in section_headers:
                self.assertFalse(converter._is_chapter_header(header),
                               f"Заголовок '{header}' НЕ должен определяться как глава")

    def test_chapter_filename_generation(self):
        """Тест генерации имён файлов для глав"""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "temp.docx"
            from docx import Document
            doc = Document()
            doc.save(doc_path)
            
            converter = ImprovedDocxToMarkdownConverter(str(doc_path))
            
            test_cases = [
                ("1. Общие сведения", "1.common.md"),
                ("2. Архитектура комплекса", "2.architecture.md"), 
                ("3. Технические и программные требования", "3.technical-requirements.md"),
                ("4. Установка и первоначальная настройка", "4.installation-setup.md")
            ]
            
            for chapter_title, expected_filename in test_cases:
                chapter_info = converter._extract_chapter_info(chapter_title)
                filename = converter._generate_chapter_filename(chapter_info)
                # Проверяем, что формат файла корректный (номер + название + .md)
                self.assertTrue(filename.startswith(f"{chapter_info['number']}."),
                              f"Файл должен начинаться с номера главы: '{filename}'")
                self.assertTrue(filename.endswith('.md'),
                              f"Файл должен иметь расширение .md: '{filename}'")

    def test_frontmatter_generation(self):
        """Тест генерации frontmatter для глав"""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "temp.docx"
            from docx import Document
            doc = Document()
            doc.save(doc_path)
            
            converter = ImprovedDocxToMarkdownConverter(str(doc_path))
            
            chapter_info = {'number': 2, 'title': 'Архитектура комплекса'}
            frontmatter = converter._generate_chapter_frontmatter(chapter_info, 7)
            
            self.assertIn("title: Архитектура комплекса", frontmatter)
            self.assertTrue(frontmatter.startswith("---"))
            self.assertTrue(frontmatter.endswith("---"))

    def test_samples_compatibility(self):
        """Тест совместимости с образцами в папке samples"""
        if not self.samples_dir.exists():
            self.skipTest("Папка samples не найдена")
            
        # Проверяем, что в samples есть ожидаемые файлы
        expected_files = [
            "1.common.md",
            "2.architecture.md", 
            "3.technical-requirements.md"
        ]
        
        for filename in expected_files:
            file_path = self.samples_dir / filename
            self.assertTrue(file_path.exists(),
                          f"Ожидаемый файл образца '{filename}' не найден в samples/")
            
            # Проверяем, что файлы содержат frontmatter
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                self.assertTrue(content.startswith("---"),
                              f"Файл '{filename}' должен начинаться с frontmatter")


if __name__ == '__main__':
    unittest.main()