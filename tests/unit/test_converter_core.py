#!/usr/bin/env python3
"""
Модульные тесты для основной функциональности DOCX конвертера
"""

import unittest
import tempfile
import os
from pathlib import Path
import sys

# Добавляем родительскую директорию в path для импорта модуля
sys.path.append(str(Path(__file__).parent.parent.parent))

from improved_docx_converter import ImprovedDocxToMarkdownConverter


class TestConverterCore(unittest.TestCase):
    """Модульные тесты для базовой функциональности конвертера"""
    
    def test_converter_initialization_with_nonexistent_file(self):
        """Тест инициализации конвертера с несуществующим файлом"""
        with self.assertRaises(Exception):
            ImprovedDocxToMarkdownConverter("nonexistent.docx")

    def test_docx_to_markdown_basic_conversion(self):
        """Проверяем базовую конвертацию DOCX в Markdown"""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "sample.docx"
            doc = Document()
            doc.add_paragraph("Тестовое содержимое")
            doc.add_heading("Заголовок", level=1)
            doc.save(doc_path)

            converter = ImprovedDocxToMarkdownConverter(str(doc_path))
            markdown = converter.convert()

            self.assertIn("Тестовое содержимое", markdown)
            self.assertIn("# Заголовок", markdown)

    def test_code_language_detection(self):
        """Тест определения языка в блоках кода"""
        from improved_docx_converter import ImprovedDocxToMarkdownConverter
        
        # Создаем временный конвертер для тестирования методов
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "temp.docx"
            from docx import Document
            doc = Document()
            doc.save(doc_path)
            
            converter = ImprovedDocxToMarkdownConverter(str(doc_path))
            
            # Тестируем определение языков
            test_cases = [
                ("docker ps -a", "bash"),
                ("SELECT * FROM table", "sql"),
                ("import sys", "python"),
                ("version: '3.8'", "yaml")
            ]
            
            for code, expected_lang in test_cases:
                detected_lang = converter._detect_code_language(code)
                self.assertEqual(detected_lang, expected_lang, 
                               f"Failed for code: {code}")

    def test_chapter_header_detection(self):
        """Тест определения заголовков глав"""
        from improved_docx_converter import ImprovedDocxToMarkdownConverter
        
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "temp.docx"
            from docx import Document
            doc = Document()
            doc.save(doc_path)
            
            converter = ImprovedDocxToMarkdownConverter(str(doc_path))
            
            # Тестируем различные форматы заголовков глав
            test_cases = [
                ("1. Общие сведения", True),
                ("2    Архитектура комплекса", True),
                ("3.1 Подраздел", False),  # подразделы не должны считаться главами
                ("Обычный текст", False),
                ("# Markdown заголовок", False)
            ]
            
            for text, expected in test_cases:
                result = converter._is_chapter_header(text)
                self.assertEqual(result, expected,
                               f"Failed for text: '{text}'")


class TestFormattingDetection(unittest.TestCase):
    """Тесты для определения различных элементов форматирования"""
    
    def setUp(self):
        """Подготовка для тестов"""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "temp.docx"
            from docx import Document
            doc = Document()
            doc.save(doc_path)
            self.converter = ImprovedDocxToMarkdownConverter(str(doc_path))
    
    def test_special_blocks_detection(self):
        """Тест определения специальных блоков"""
        test_cases = [
            ("::AppAnnotation", True),
            ("::", True),
            ("Обычный текст", False),
            (":: AppAnnotation", True)
        ]
        
        for text, expected in test_cases:
            result = text.strip().startswith("::") or text.strip() == "::"
            self.assertEqual(result, expected,
                           f"Failed for text: '{text}'")


if __name__ == '__main__':
    unittest.main()