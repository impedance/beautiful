#!/usr/bin/env python3
"""
Интеграционные тесты для обработки реальных DOCX файлов
"""

import unittest
import tempfile
import os
import shutil
from pathlib import Path
import sys

# Добавляем родительскую директорию в path для импорта модуля
sys.path.append(str(Path(__file__).parent.parent.parent))

from improved_docx_converter import ImprovedDocxToMarkdownConverter


class TestDocxProcessing(unittest.TestCase):
    """Интеграционные тесты для обработки DOCX файлов"""
    
    def setUp(self):
        """Подготовка к тестам"""
        self.temp_dir = Path(tempfile.mkdtemp())
        self.test_docx = Path("dev.docx")
        
    def tearDown(self):
        """Очистка после тестов"""
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)

    def test_real_docx_conversion(self):
        """Тест конвертации реального DOCX файла"""
        if not self.test_docx.exists():
            self.skipTest("Тестовый DOCX файл dev.docx не найден")
            
        converter = ImprovedDocxToMarkdownConverter(str(self.test_docx))
        markdown = converter.convert()
        
        # Проверяем, что результат содержит ожидаемые элементы
        self.assertIsInstance(markdown, str)
        self.assertGreater(len(markdown), 100)  # Должен содержать существенный контент
        
        # Проверяем наличие основных элементов документа
        self.assertIn("Общие сведения", markdown)
        
    def test_chapter_splitting_with_real_docx(self):
        """Тест разбиения реального DOCX на главы"""
        if not self.test_docx.exists():
            self.skipTest("Тестовый DOCX файл dev.docx не найден")
            
        converter = ImprovedDocxToMarkdownConverter(str(self.test_docx), split_chapters=True)
        
        # Создаем главы в временной директории
        os.chdir(self.temp_dir)
        chapters = converter.convert_to_chapters()
        
        # Проверяем, что главы были созданы
        self.assertGreater(len(chapters), 0, "Должны быть созданы главы")
        
        # Проверяем, что файлы глав существуют  
        for chapter_filename in chapters:
            chapter_file = Path("chapters") / chapter_filename
            self.assertTrue(chapter_file.exists(),
                          f"Файл главы '{chapter_filename}' должен существовать")
            
            # Проверяем, что файл содержит контент
            with open(chapter_file, 'r', encoding='utf-8') as f:
                content = f.read()
                self.assertGreater(len(content), 10,
                                 f"Файл '{chapter_filename}' должен содержать контент")

    def test_regex_patterns_with_real_content(self):
        """Тест regex паттернов на реальном содержимом"""
        if not self.test_docx.exists():
            self.skipTest("Тестовый DOCX файл dev.docx не найден")
            
        converter = ImprovedDocxToMarkdownConverter(str(self.test_docx))
        
        # Тестируем определение языков в блоках кода
        test_cases = [
            ("docker ps -a", "bash"),
            ("SELECT * FROM table", "sql"),  
            ("version: '3.8'", "yaml")
        ]
        
        for code, expected_lang in test_cases:
            detected_lang = converter._detect_code_language(code)
            self.assertEqual(detected_lang, expected_lang,
                           f"Неправильное определение языка для: {code}")


class TestTableProcessing(unittest.TestCase):
    """Тесты для обработки таблиц в DOCX"""
    
    def test_table_formatting_detection(self):
        """Тест определения и форматирования таблиц"""
        from docx import Document
        from docx.shared import Inches
        
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "table_test.docx"
            doc = Document()
            
            # Создаем таблицу для тестирования
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Заголовок 1"
            table.cell(0, 1).text = "Заголовок 2"
            table.cell(1, 0).text = "Данные 1"
            table.cell(1, 1).text = "Данные 2"
            
            doc.save(doc_path)
            
            converter = ImprovedDocxToMarkdownConverter(str(doc_path))
            markdown = converter.convert()
            
            # Проверяем, что таблица правильно конвертирована
            self.assertIn("| Заголовок 1 | Заголовок 2 |", markdown)
            self.assertIn("| Данные 1 | Данные 2 |", markdown)
            self.assertIn("|", markdown)  # Должны быть разделители таблицы

    def test_table_cell_formatting(self):
        """Тест форматирования ячеек таблицы"""
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "cell_test.docx"
            doc = Document()
            
            # Создаем таблицу с переносами строк в ячейках
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            p1 = cell.paragraphs[0]
            p1.text = "Первая строка"
            p2 = cell.add_paragraph("Вторая строка")
            
            doc.save(doc_path)
            
            converter = ImprovedDocxToMarkdownConverter(str(doc_path))
            markdown = converter.convert()
            
            # Проверяем, что переносы строк в ячейках обработаны
            self.assertIn("Первая строка", markdown)
            self.assertIn("Вторая строка", markdown)


if __name__ == '__main__':
    unittest.main()