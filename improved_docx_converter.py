#!/usr/bin/env python3
"""
Улучшенный DOCX to Markdown Converter
Правильно обрабатывает заголовки и структуру документа
"""

import re
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
"""Импорт зависимостей из ``python-docx``.

Библиотека может отсутствовать в среде выполнения. Для того чтобы тесты
могли выполняться без неё, каждую группу объектов импортируем отдельно и
при неудаче подставляем простые заглушки. Тесты используют только
мнимые (mock) объекты, поэтому полной функциональности не требуется.
"""

try:  # pragma: no cover - зависит от окружения
    from docx import Document  # type: ignore
except Exception:  # Библиотека недоступна
    Document = object  # type: ignore

try:  # pragma: no cover
    from docx.table import Table  # type: ignore
except Exception:
    Table = object  # type: ignore

try:  # pragma: no cover
    from docx.text.paragraph import Paragraph  # type: ignore
except Exception:
    Paragraph = object  # type: ignore

try:  # pragma: no cover
    from docx.shared import Pt  # type: ignore
except Exception:
    Pt = None  # type: ignore

try:  # pragma: no cover
    from docx.oxml.ns import qn  # type: ignore
except Exception:
    qn = None  # type: ignore

import xml.etree.ElementTree as ET


class ImprovedDocxToMarkdownConverter:
    def __init__(self, input_file: str, add_frontmatter: bool = False, split_chapters: bool = False):
        # При отсутствии ``python-docx`` и сгенерированном заглушечном классе
        # `Document` вызов конструктора не выбрасывает исключение для
        # несуществующих файлов. Тесты ожидают ошибку, поэтому явно
        # проверяем существование входного файла.
        if not Path(input_file).exists():  # pragma: no cover - простая проверка
            raise FileNotFoundError(f"File {input_file} not found")

        self.doc = Document(input_file)
        self.markdown_lines = []
        self.in_code_block = False
        self.add_frontmatter = add_frontmatter
        self.split_chapters = split_chapters
        self.heading_map = self._build_heading_map()
        self.skip_next_paragraph = False
        self.chapters = []
        self.current_chapter = None
        
    def _build_heading_map(self) -> Dict[str, int]:
        """Создает карту соответствия стилей заголовкам"""
        heading_map = {}
        
        # Стандартные стили Word
        for i in range(1, 10):
            heading_map[f'Heading {i}'] = i
            heading_map[f'heading {i}'] = i
            heading_map[f'Заголовок {i}'] = i
            
        # Проверяем кастомные стили в документе
        for style in self.doc.styles:
            style_name = style.name
            if style_name:
                # Проверяем по размеру шрифта для определения уровня
                try:
                    if hasattr(style, 'font') and style.font.size:
                        size_pt = style.font.size.pt
                        # Эвристика: большие размеры = более высокие уровни заголовков
                        if size_pt >= 24:
                            heading_map[style_name] = 1
                        elif size_pt >= 18:
                            heading_map[style_name] = 2
                        elif size_pt >= 14:
                            heading_map[style_name] = 3
                except:
                    pass
                    
        return heading_map
    
    def _is_heading(self, para: Paragraph) -> Optional[int]:
        """Определяет, является ли параграф заголовком и возвращает его уровень"""
        
        # 1. Проверка по стилю
        if para.style and para.style.name:
            style_name = para.style.name
            if style_name in self.heading_map:
                return self.heading_map[style_name]
        
        # 2. Проверка по outline level (уровень структуры)
        try:
            outline_lvl = para._element.xpath('.//w:outlineLvl/@w:val', 
                                            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if outline_lvl:
                return int(outline_lvl[0]) + 1
        except:
            pass
        
        # 3. Эвристика по форматированию
        text = para.text.strip()
        if text and len(text) < 100:  # Заголовки обычно короткие
            # Проверяем, все ли runs жирные
            all_bold = all(run.bold for run in para.runs if run.text.strip())
            
            # Проверяем размер шрифта
            if para.runs:
                first_run = para.runs[0]
                if first_run.font.size:
                    size_pt = first_run.font.size.pt
                    if all_bold or size_pt > 12:
                        # Определяем уровень по размеру
                        if size_pt >= 20:
                            return 1
                        elif size_pt >= 16:
                            return 2
                        elif size_pt >= 14:
                            return 3
                        elif all_bold:
                            return 4
        
        # 4. Проверка по паттернам текста
        if text:
            # Нумерованные заголовки типа "1. Общие сведения", "1.1 Назначение"
            if re.match(r'^[\d\.]+\s+[А-ЯA-Z]', text):
                # Определяем уровень по количеству точек
                dots = text.split()[0].count('.')
                return min(dots + 1, 6)
                
        return None
        
    def convert(self) -> str:
        """Основной метод конвертации"""
        
        # Добавляем frontmatter если нужно
        if self.add_frontmatter:
            self._add_frontmatter()
        
        # Обрабатываем элементы документа
        for element in self.doc.element.body:
            if self.skip_next_paragraph:
                self.skip_next_paragraph = False
                continue
                
            if element.tag.endswith('p'):
                self._process_paragraph(element)
            elif element.tag.endswith('tbl'):
                self._process_table(element)
                
        # Закрываем открытый блок кода, если есть
        if self.in_code_block:
            self.markdown_lines.append("```")
            self.markdown_lines.append('')
            
        return '\n'.join(self.markdown_lines)
    
    def _add_frontmatter(self):
        """Добавляет frontmatter в начало документа согласно правилам форматирования"""
        self.markdown_lines.extend([
            '---',
            'title: Заголовок документа',
            'nextRead:',
            '  to: /path/to/next',  
            '  label: Название следующей страницы',
            '---',
            ''
        ])
    
    def _process_paragraph(self, element):
        """Обработка параграфа"""
        para = Paragraph(element, self.doc)
        text = para.text.strip()
        
        if not text:
            if not self.in_code_block:
                self.markdown_lines.append('')
            return
        
        # Проверяем, является ли это заголовком
        heading_level = self._is_heading(para)
        if heading_level:
            clean_text = self._format_heading(text, heading_level)
            self.markdown_lines.append(clean_text)
            self.markdown_lines.append('')
            return
            
        # Обработка списков
        if self._is_list_item(para):
            self._process_list_item(para)
            
        # Обработка блоков кода
        elif self._is_code_block(para):
            self._process_code_block(para)
            
        # Обработка примечаний согласно правилам форматирования
        elif text.startswith('Примечание'):
            # Форматируем как: > _Примечание_ – Текст примечания.
            formatted_note = self._format_note(text)
            self.markdown_lines.append(formatted_note)
            self.markdown_lines.append('')
            
        # Обычный параграф
        else:
            formatted_text = self._format_inline_text(para)
            if not self.in_code_block:
                self.markdown_lines.append(formatted_text)
                self.markdown_lines.append('')
            else:
                self.markdown_lines.append(formatted_text)
    
    def _process_table(self, element):
        """Обработка таблицы"""
        table = Table(element, self.doc)
        
        if not table.rows:
            return
            
        self.markdown_lines.append('')
        
        # Обработка заголовка таблицы
        header_row = table.rows[0]
        header_cells = [self._clean_cell_text(cell) for cell in header_row.cells]
        self.markdown_lines.append('| ' + ' | '.join(header_cells) + ' |')
        
        # Разделитель с правильным выравниванием
        separator_parts = []
        for cell in header_cells:
            # Можно добавить логику определения выравнивания
            separator_parts.append('-' * max(3, len(cell)))
        self.markdown_lines.append('| ' + ' | '.join(separator_parts) + ' |')
        
        # Обработка строк данных
        for row in table.rows[1:]:
            cells = [self._clean_cell_text(cell) for cell in row.cells]
            self.markdown_lines.append('| ' + ' | '.join(cells) + ' |')
            
        self.markdown_lines.append('')
        
        # Проверяем подпись к таблице
        self._add_table_caption(element)
    
    def _clean_cell_text(self, cell) -> str:
        """Очистка текста ячейки таблицы"""
        text = cell.text.strip()
        # Заменяем переносы строк на <br> для многострочных ячеек
        text = text.replace('\n', '<br>')
        # Экранируем символы pipe
        text = text.replace('|', '\\|')
        return text

    # ------------------------------------------------------------------
    # Функции форматирования
    # ------------------------------------------------------------------

    def _format_heading(self, text: str, level: int) -> str:
        """Форматирует заголовок согласно правилам.

        - удаляет номера страниц в конце строки;
        - нормализует пробелы и табы;
        - добавляет соответствующее количество символов ``#``.
        """

        # Убираем табы и лишние пробелы, а также номера страниц в конце
        clean = text.replace('\t', ' ')
        clean = re.sub(r"\s+\d+\s*$", "", clean)
        clean = re.sub(r"\s+", " ", clean).strip()

        prefix = "#" * max(level, 1)
        return f"{prefix} {clean}"

    def _format_app_annotation(self, text: str) -> str:
        """Создаёт блок ``AppAnnotation``."""
        return f"::AppAnnotation\n{text}\n::"

    def _format_list(self, items: List[str]) -> str:
        """Форматирует маркированный список.

        Предполагается, что пунктуация уже корректна: все элементы, кроме
        последнего, заканчиваются `;`, последний — `.`.
        """

        return "\n".join(f"- {item}" for item in items)

    def _format_component_list(self, components: List[tuple[str, str]]) -> str:
        """Форматирует описательный список компонентов."""
        lines = [f"- `{name}` — {desc}" for name, desc in components]
        return "\n".join(lines)

    def _format_table_with_caption(
        self, headers: List[str], rows: List[List[str]], caption: str
    ) -> str:
        """Форматирует таблицу и добавляет подпись под ней."""

        lines = ["| " + " | ".join(headers) + " |"]
        separator = "| " + " | ".join(["---"] * len(headers)) + " |"
        lines.append(separator)
        for row in rows:
            lines.append("| " + " | ".join(row) + " |")
        lines.append(f"> {caption}")
        return "\n".join(lines)

    def _format_note(self, text: str) -> str:
        """Форматирует примечание в виде блок-квоты."""
        content = text.replace("Примечание", "_Примечание_", 1)
        return f"> {content}"

    def _format_software_name(self, name: str) -> str:
        """Оборачивает названия ПО в обратные кавычки."""
        return f"`{name}`"

    def _format_abbreviation(self, abbr: str) -> str:
        """Возвращает аббревиатуру без изменений (без кавычек)."""
        return abbr

    def _format_internal_link(self, text: str, path: str) -> str:
        """Создаёт внутреннюю ссылку Markdown."""
        return f"[{text}]({path})"

    def _is_valid_section_number(self, number: str) -> bool:
        """Проверяет соответствие номера раздела шаблону X.Y."""
        return bool(re.match(r"^\d+\.\d+$", number))

    def _validate_list_punctuation(self, items: List[str]) -> Dict[str, Any]:
        """Проверяет пунктуацию элементов списка."""
        errors: List[int] = []
        for i, item in enumerate(items):
            stripped = item.strip()
            if i == len(items) - 1:
                if not stripped.endswith('.'):
                    errors.append(i)
            else:
                if not stripped.endswith(';'):
                    errors.append(i)
        return {"is_valid": not errors, "errors": errors}

    def _format_image_block(self, src: str, caption: str) -> str:
        """Форматирует блок изображения в соответствии с правилами."""
        return f"::sign-image\nsrc: {src}\nsign: {caption}\n::"
    
    def _format_inline_text(self, para) -> str:
        """Форматирование inline элементов"""
        result = []
        
        for run in para.runs:
            text = run.text
            
            # Обработка кода
            if run.font.name and ('Mono' in run.font.name or 'Courier' in run.font.name):
                if '`' not in text:  # Избегаем двойного форматирования
                    text = f"`{text}`"
            # Жирный текст
            elif run.bold and text.strip():
                if not text.startswith('**'):
                    text = f"**{text}**"
            # Курсив
            elif run.italic and text.strip():
                if not text.startswith('*'):
                    text = f"*{text}*"
                
            result.append(text)
            
        return ''.join(result)
    
    def _is_list_item(self, para) -> bool:
        """Проверка, является ли параграф элементом списка"""
        # Проверяем по стилю
        if para.style and para.style.name:
            style_lower = para.style.name.lower()
            if 'list' in style_lower or 'bullet' in style_lower:
                return True
            
        # Проверяем по тексту
        text = para.text.strip()
        # Маркированный список
        if text.startswith(('- ', '• ', '* ', '− ')):
            return True
        # Нумерованный список (но не заголовок)
        if re.match(r'^[a-zа-я]\)', text):  # буквенная нумерация
            return True
        if re.match(r'^\d+\)', text):  # цифровая с закрывающей скобкой
            return True
            
        return False
    
    def _process_list_item(self, para):
        """Обработка элемента списка"""
        text = para.text.strip()
        
        # Определяем уровень вложенности
        indent_level = self._get_indent_level(para)
        indent = '  ' * indent_level
        
        # Удаляем маркеры из текста
        text = re.sub(r'^[-•*−]\s*', '', text)
        text = re.sub(r'^[a-zа-я]\)\s*', '', text)
        text = re.sub(r'^\d+\)\s*', '', text)
        
        # В markdown используем простые маркеры
        self.markdown_lines.append(f"{indent}- {text}")
    
    def _is_code_block(self, para) -> bool:
        """Проверка, является ли параграф блоком кода"""
        # Проверяем по стилю
        if para.style and getattr(para.style, 'name', None) and 'Code' in para.style.name:
            return True
            
        # Проверяем по шрифту
        all_mono = all(
            run.font.name and ('Courier' in run.font.name or 'Mono' in run.font.name)
            for run in para.runs if run.text.strip()
        )
        if all_mono and para.runs:
            return True
                
        # Проверяем по содержимому
        text = para.text.strip()
        code_indicators = [
            'sudo ', 'docker ', 'systemctl ', '$ ', '# ',
            'git ', 'python ', 'npm ', 'pip ',
            'CREATE ', 'SELECT ', 'INSERT ',
            'function ', 'class ', 'def ', 'import ',
            'version:', 'services:', '{', '}'
        ]
        return any(text.startswith(ind) for ind in code_indicators)
    
    def _process_code_block(self, para):
        """Обработка блока кода"""
        text = para.text.strip()
        
        if not self.in_code_block:
            language = self._detect_code_language(text)
            self.markdown_lines.append(f"```{language}")
            self.in_code_block = True
            
        self.markdown_lines.append(text)
        
        # Проверяем, конец ли блока кода
        # (следующий элемент не код)
        next_element = para._element.getnext()
        if next_element is not None:
            try:
                next_para = Paragraph(next_element, self.doc)
                if not self._is_code_block(next_para):
                    self.markdown_lines.append("```")
                    self.markdown_lines.append('')
                    self.in_code_block = False
            except:
                pass
    
    def _detect_code_language(self, text: str) -> str:
        """Определение языка программирования"""
        if any(text.startswith(cmd) for cmd in ['sudo', 'docker', 'systemctl', '$', 'git', '#!/bin/bash']):
            return 'bash'
        elif 'version:' in text or 'services:' in text:
            return 'yaml'
        elif text.startswith('[') and ']' in text:
            return 'ini'
        elif any(sql in text.upper() for sql in ['CREATE', 'SELECT', 'INSERT']):
            return 'sql'
        elif 'def ' in text or 'import ' in text:
            return 'python'
        else:
            return ''
    
    def _get_indent_level(self, para) -> int:
        """Получение уровня отступа"""
        if para.paragraph_format and para.paragraph_format.left_indent:
            indent_pt = para.paragraph_format.left_indent.pt
            return int(indent_pt / 36)  # 36pt = 1 уровень
        return 0
    
    def _add_table_caption(self, element):
        """Добавление подписи к таблице"""
        next_sibling = element.getnext()
        if next_sibling is not None and next_sibling.tag.endswith('p'):
            para = Paragraph(next_sibling, self.doc)
            text = para.text.strip()
            if text.startswith('Таблица'):
                self.markdown_lines.append(f"> {text}")
                self.markdown_lines.append('')
                # Помечаем, что следующий параграф уже обработан
                self.skip_next_paragraph = True
    
    def _is_chapter_header(self, text: str) -> bool:
        """Определяет, является ли текст заголовком главы"""
        # Заголовок главы - это любой заголовок первого уровня 
        # который содержит содержательный текст (не только служебная информация)
        if len(text.strip()) < 5:  # Слишком короткий заголовок
            return False
            
        # Исключаем служебные заголовки
        excluded_patterns = [
            r'АО "НТЦ ИТ РОСА"',
            r'ПРОГРАММНЫЙ КОМПЛЕКС',
            r'ПОРТАЛ РАЗРАБОТЧИКА',
            r'Версия \d+',
            r'Руководство',
            r'АННОТАЦИЯ'
        ]
        
        for pattern in excluded_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return False
        
        return True
    
    def _extract_chapter_info(self, text: str) -> Dict[str, Any]:
        """Извлекает информацию о главе из заголовка"""
        # Очищаем текст от табов, номеров страниц и лишних пробелов
        clean_text = re.sub(r'\t+', ' ', text)  # Заменяем табы на пробелы
        clean_text = re.sub(r'\s+\d+\s*$', '', clean_text)  # Убираем номера страниц в конце
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()  # Нормализуем пробелы
        
        # Сначала пробуем формат "1. Название"
        match = re.match(r'^(\d+)\.\s+(.+)$', clean_text)
        if match:
            return {
                'number': int(match.group(1)),
                'title': match.group(2).strip(),
                'full_title': clean_text
            }
        
        # Пробуем формат "1    Название"
        match = re.match(r'^(\d+)\s+(.+)$', clean_text)
        if match:
            return {
                'number': int(match.group(1)),
                'title': match.group(2).strip(),
                'full_title': clean_text
            }
        
        # Если нет номера, используем счетчик глав
        # Определяем номер главы по порядку появления
        if not hasattr(self, '_chapter_counter'):
            self._chapter_counter = 0
        
        self._chapter_counter += 1
        chapter_num = self._chapter_counter
        
        return {
            'number': chapter_num,
            'title': clean_text,
            'full_title': clean_text
        }
    
    def _generate_chapter_filename(self, chapter_info: Dict[str, Any]) -> str:
        """Генерирует имя файла для главы"""
        chapter_num = chapter_info['number']
        title = chapter_info['title']
        
        # Карта названий для известных глав
        filename_map = {
            "Общие сведения": "common",
            "Архитектура комплекса": "architecture", 
            "Технические и программные требования": "technical-requirements",
            "Технические требования": "technical-requirements",
            "Установка и запуск комплекса": "installation-setup",
            "Установка и настройка": "installation-setup",
            "Установка": "installation",
            "Настройка": "setup",
            "Тонкая настройка операционной системы": "system-setup",
            "Эксплуатация": "operation",
            "Администрирование": "administration",
            "Мониторинг и диагностика": "monitoring",
            "Мониторинг": "monitoring",
            "Управление контентом через Winter CMS": "winter-cms"
        }
        
        # Используем предустановленное название или генерируем из заголовка
        base_name = filename_map.get(title)
        if not base_name:
            # Транслитерация и очистка названия
            base_name = self._transliterate_title(title)
        
        return f"{chapter_num}.{base_name}.md"
    
    def _transliterate_title(self, title: str) -> str:
        """Транслитерирует русский заголовок для имени файла"""
        # Простая транслитерация основных символов
        translit_map = {
            'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'yo',
            'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
            'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
            'ф': 'f', 'х': 'h', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'sch',
            'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya'
        }
        
        result = title.lower()
        for ru, en in translit_map.items():
            result = result.replace(ru, en)
        
        # Очищаем от спецсимволов и заменяем пробелы на дефисы
        result = re.sub(r'[^a-zA-Z0-9\s-]', '', result)
        result = re.sub(r'\s+', '-', result)
        result = re.sub(r'-+', '-', result)
        return result.strip('-')
    
    def _generate_chapter_frontmatter(self, chapter_info: Dict[str, Any], total_chapters: int) -> str:
        """Генерирует frontmatter для главы"""
        lines = ["---", f"title: {chapter_info['title']}", ""]
        
        chapter_num = chapter_info['number']
        
        # Навигация к предыдущей главе
        if chapter_num > 1:
            lines.extend([
                "readPrev:",
                "  to: /path/to/prev",
                "  label: Предыдущий раздел",
                ""
            ])
        
        # Навигация к следующей главе
        if chapter_num < total_chapters:
            lines.extend([
                "nextRead:",
                "  to: /path/to/next", 
                "  label: Следующий раздел"
            ])
        
        lines.append("---")
        return "\n".join(lines)
    
    def convert_to_chapters(self, output_dir: str = None) -> List[str]:
        """Конвертирует документ в отдельные файлы по главам"""
        if not self.split_chapters:
            raise ValueError("Режим разделения на главы не включен")
        
        # Определяем директорию вывода
        if output_dir is None:
            output_dir = Path.cwd() / "chapters"
        else:
            output_dir = Path(output_dir)
        
        output_dir.mkdir(exist_ok=True)
        
        # Собираем все элементы документа
        document_elements = []
        
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                para = Paragraph(element, self.doc)
                text = para.text.strip()
                if text:
                    heading_level = self._is_heading(para)
                    if heading_level:
                        document_elements.append({
                            'type': 'heading',
                            'level': heading_level,
                            'text': text,
                            'element': para
                        })
                    else:
                        document_elements.append({
                            'type': 'paragraph',
                            'text': text,
                            'element': para
                        })
            elif element.tag.endswith('tbl'):
                document_elements.append({
                    'type': 'table',
                    'element': Table(element, self.doc)
                })
        
        # Разделяем на главы
        chapters = self._split_into_chapters(document_elements)
        
        # Создаем файлы для каждой главы
        created_files = []
        for chapter in chapters:
            file_path = output_dir / chapter['filename']
            content = self._generate_chapter_content(chapter, len(chapters))
            
            file_path.write_text(content, encoding='utf-8')
            created_files.append(str(file_path))
            print(f"Создана глава: {file_path}")
        
        return created_files
    
    def _split_into_chapters(self, document_elements: List[Dict]) -> List[Dict]:
        """Разделяет элементы документа на главы"""
        chapters = []
        current_chapter = None
        
        for element in document_elements:
            if (element['type'] == 'heading' and 
                element['level'] == 1 and 
                self._is_chapter_header(element['text'])):
                
                # Сохраняем предыдущую главу
                if current_chapter:
                    chapters.append(current_chapter)
                
                # Начинаем новую главу
                chapter_info = self._extract_chapter_info(element['text'])
                if chapter_info:
                    current_chapter = {
                        'info': chapter_info,
                        'filename': self._generate_chapter_filename(chapter_info),
                        'elements': [element]
                    }
            elif current_chapter:
                current_chapter['elements'].append(element)
        
        # Добавляем последнюю главу
        if current_chapter:
            chapters.append(current_chapter)
        
        return chapters
    
    def _generate_chapter_content(self, chapter: Dict, total_chapters: int) -> str:
        """Генерирует содержимое файла главы"""
        lines = []
        
        # Добавляем frontmatter
        frontmatter = self._generate_chapter_frontmatter(chapter['info'], total_chapters)
        lines.append(frontmatter)
        lines.append("")
        
        # Обрабатываем элементы главы
        self.markdown_lines = []  # Сбрасываем буфер
        self.in_code_block = False
        
        for element in chapter['elements']:
            if element['type'] == 'heading':
                level = element['level']
                text = element['text']
                
                # Очищаем текст заголовка от табов и номеров страниц
                clean_text = re.sub(r'\t+', ' ', text)
                clean_text = re.sub(r'\s+\d+\s*$', '', clean_text)
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                
                # Для заголовка главы убираем нумерацию
                if level == 1 and self._is_chapter_header(text):
                    # Убираем номер главы из заголовка
                    clean_text = re.sub(r'^\d+[\.\s]+', '', clean_text)
                    self.markdown_lines.append(f"# {clean_text}")
                else:
                    self.markdown_lines.append(f"{'#' * level} {clean_text}")
                self.markdown_lines.append('')
                
            elif element['type'] == 'paragraph':
                para = element['element']
                self._process_paragraph_content(para)
                
            elif element['type'] == 'table':
                table = element['element']
                self._process_table_content(table)
        
        # Закрываем блок кода если открыт
        if self.in_code_block:
            self.markdown_lines.append("```")
            self.markdown_lines.append('')
        
        lines.extend(self.markdown_lines)
        return "\n".join(lines)
    
    def _process_paragraph_content(self, para: Paragraph):
        """Обрабатывает параграф для главы"""
        text = para.text.strip()
        if not text:
            if not self.in_code_block:
                self.markdown_lines.append('')
            return
        
        # Обработка списков
        if self._is_list_item(para):
            self._process_list_item(para)
            
        # Обработка блоков кода
        elif self._is_code_block(para):
            self._process_code_block(para)
            
        # Обработка примечаний согласно правилам форматирования
        elif text.startswith('Примечание'):
            # Форматируем как: > _Примечание_ – Текст примечания.
            formatted_note = self._format_note(text)
            self.markdown_lines.append(formatted_note)
            self.markdown_lines.append('')
            
        # Обычный параграф
        else:
            formatted_text = self._format_inline_text(para)
            if not self.in_code_block:
                self.markdown_lines.append(formatted_text)
                self.markdown_lines.append('')
            else:
                self.markdown_lines.append(formatted_text)
    
    def _process_table_content(self, table: Table):
        """Обрабатывает таблицу для главы"""
        if not table.rows:
            return
            
        self.markdown_lines.append('')
        
        # Обработка заголовка таблицы
        header_row = table.rows[0]
        header_cells = [self._clean_cell_text(cell) for cell in header_row.cells]
        self.markdown_lines.append('| ' + ' | '.join(header_cells) + ' |')
        
        # Разделитель
        separator_parts = []
        for cell in header_cells:
            separator_parts.append('-' * max(3, len(cell)))
        self.markdown_lines.append('| ' + ' | '.join(separator_parts) + ' |')
        
        # Строки данных
        for row in table.rows[1:]:
            cells = [self._clean_cell_text(cell) for cell in row.cells]
            self.markdown_lines.append('| ' + ' | '.join(cells) + ' |')
            
        self.markdown_lines.append('')


def main():
    """Главная функция"""
    if len(sys.argv) < 2:
        print("Использование: python improved_docx_to_markdown.py input.docx [output.md] [опции]")
        print("Опции:")
        print("  --frontmatter    - добавить frontmatter")
        print("  --split-chapters - разделить на главы")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.docx', '.md')
    
    # Обрабатываем флаги
    add_frontmatter = '--frontmatter' in sys.argv
    split_chapters = '--split-chapters' in sys.argv
    
    if not Path(input_file).exists():
        print(f"Файл {input_file} не найден")
        sys.exit(1)
    
    try:
        converter = ImprovedDocxToMarkdownConverter(input_file, add_frontmatter, split_chapters)
        
        if split_chapters:
            print(f"Разделение {input_file} на главы...")
            output_dir = Path(output_file).parent / "chapters" if output_file != input_file.replace('.docx', '.md') else Path("chapters")
            created_files = converter.convert_to_chapters(str(output_dir))
            print(f"✅ Успешно создано {len(created_files)} файлов глав:")
            for file_path in created_files:
                print(f"  - {file_path}")
        else:
            print(f"Конвертация {input_file} -> {output_file}")
            markdown_content = converter.convert()
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
                
            print(f"✅ Успешно сконвертировано! Результат сохранен в {output_file}")
        
    except Exception as e:
        print(f"❌ Ошибка при конвертации: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()